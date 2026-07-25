"""
Raw text extraction from an uploaded CV file. Supports the three formats
a real candidate is realistically going to upload: PDF, DOCX, and plain
text. This module only extracts text -- it has no opinion about what the
text means; that's cv_parser.py's job.

Deliberately does not attempt OCR for scanned/image-only PDFs -- pypdf
will return an empty or near-empty string for those, which
extract_text_from_file() treats as an extraction failure (see
MIN_EXTRACTED_CHARS below) rather than silently handing the LLM an
empty prompt and getting back a hallucinated profile. Real OCR is a
known follow-up, not attempted here.
"""
from io import BytesIO

import docx
from pypdf import PdfReader

# Below this length, treat extraction as having failed rather than
# passing near-empty text to the LLM -- a genuinely short CV is still
# almost always well over this (a single line of text on a page is
# already ~40-60 chars); this is really a scanned-PDF/corrupt-file
# tripwire, not a real length judgment on CV content.
MIN_EXTRACTED_CHARS = 100

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt")


class ExtractionError(Exception):
    """Raised when a file can't be read as one of the supported CV
    formats, or when extraction produces too little text to be a real
    CV (e.g. a scanned/image-only PDF with no embedded text layer).
    """


def extract_text_from_file(filename: str, content: bytes) -> str:
    """Dispatches on file extension and returns extracted plain text.

    Raises ExtractionError for an unsupported extension, a file that
    can't be parsed as its claimed format, or extracted text under
    MIN_EXTRACTED_CHARS.
    """
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        text = _extract_pdf(content)
    elif lower_name.endswith(".docx"):
        text = _extract_docx(content)
    elif lower_name.endswith(".txt"):
        text = _extract_txt(content)
    else:
        raise ExtractionError(
            f"Unsupported file type for '{filename}'. "
            f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}."
        )

    text = text.strip()
    if len(text) < MIN_EXTRACTED_CHARS:
        raise ExtractionError(
            f"Extracted only {len(text)} characters from '{filename}' -- "
            "too little to be a real CV. This usually means the PDF is a "
            "scanned image with no embedded text layer (OCR isn't "
            "supported yet), or the file is corrupt/empty."
        )
    return text


def _extract_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:
        raise ExtractionError(f"Could not open file as a PDF: {exc}") from exc

    pages_text = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:
            # A single malformed page shouldn't kill extraction for the
            # rest of an otherwise-readable document.
            continue
    return "\n".join(pages_text)


def _extract_docx(content: bytes) -> str:
    try:
        document = docx.Document(BytesIO(content))
    except Exception as exc:
        raise ExtractionError(f"Could not open file as a DOCX: {exc}") from exc

    paragraphs = [p.text for p in document.paragraphs]

    # Plenty of real CVs put content in tables (skills grids, two-column
    # layouts) rather than plain paragraphs -- skipping these would
    # silently drop real content for those documents.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def _extract_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content.decode("latin-1")
        except Exception as exc:
            raise ExtractionError(f"Could not decode text file: {exc}") from exc
